"""Crea il template DOCX aziendale pulito con placeholder, partendo dal documento di esempio.

Strategia: apriamo il doc originale, puliamo le tabelle di copertina/metadati con placeholder,
puliamo l'header, e rimuoviamo il contenuto (paragrafi e tabelle dopo le prime 2 tabelle)
preservando la sezione e il suo header/footer.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

SRC = Path(__file__).parent / "SIPAD Analisi Cr-Evolutive CONTENZIOSO-V 1.0 - 25289 - Nuovi campi 'Spese'.docx"
DST = Path(__file__).parent / "template_aziendale.docx"

doc = Document(str(SRC))

# ── 1. Pulisci tabella copertina (tabella 0) ────────────────────────────
cover = doc.tables[0]

# Riga 0: cella 0 e 1 sono merged — riscrivi
for cell_idx in range(2):
    cell = cover.cell(0, cell_idx)
    for p in cell.paragraphs:
        p.clear()
    run = cell.paragraphs[0].add_run("{{INTESTAZIONE_ENTE}}")
    run.font.size = Pt(18)
    run.bold = True

# Riga 1-4: pulisci solo colonna 1 (i valori)
cover_values = {
    1: "{{NOME_DOCUMENTO}}",
    2: "{{VERSIONE}}",
    3: "{{STATO}}",
    4: "{{DATA}}",
}
for row_idx, placeholder in cover_values.items():
    cell = cover.cell(row_idx, 1)
    for p in cell.paragraphs:
        p.clear()
    run = cell.paragraphs[0].add_run(placeholder)
    run.bold = True

# ── 2. Pulisci tabella metadati (tabella 1) ─────────────────────────────
meta = doc.tables[1]
meta_fields = {
    0: "{{CLIENTE}}",
    1: "{{PROGETTO}}",
    2: "{{DATA}}",
    3: "{{REDATTO_DA}}",
    4: "{{APPROVATO_DA}}",
    5: "{{VERIFICATO_DA}}",
    6: "{{VERSIONE}}",
    7: "{{NOME_DOCUMENTO}}",
}
for row_idx, placeholder in meta_fields.items():
    cell = meta.cell(row_idx, 1)
    for p in cell.paragraphs:
        p.clear()
    run = cell.paragraphs[0].add_run(placeholder)
    run.font.size = Pt(10)

# ── 3. Rimuovi contenuto ma preserva sezioni ────────────────────────────
body = doc.element.body
cover_tbl = cover._tbl
meta_tbl = meta._tbl

# Raccogli tutti gli elementi da rimuovere: tutto dopo meta_tbl tranne w:sectPr
found_meta = False
to_remove = []
for elem in list(body):
    if elem is meta_tbl:
        found_meta = True
        continue
    if found_meta:
        if elem.tag == qn("w:sectPr"):
            continue
        to_remove.append(elem)

for elem in to_remove:
    body.remove(elem)

# Aggiungi un paragrafo vuoto prima del sectPr come segnaposto
new_p = OxmlElement("w:p")
sect_pr = body.find(qn("w:sectPr"))
if sect_pr is not None:
    sect_pr.addprevious(new_p)
else:
    body.append(new_p)

# ── 4. Pulisci header con placeholder ───────────────────────────────────
for section in doc.sections:
    header = section.header
    if header and header.tables:
        ht = header.tables[0]
        n_cols = len(ht.columns)
        n_rows = len(ht.rows)

        # Riga 0: colonna 0 aveva il logo — pulisci (no logo)
        for p in ht.cell(0, 0).paragraphs:
            p.clear()

        if n_cols >= 4:
            for c in range(1, n_cols):
                for p in ht.cell(0, c).paragraphs:
                    p.clear()
            run = ht.cell(0, 1).paragraphs[0].add_run("{{INTESTAZIONE_ENTE}}")
            run.font.size = Pt(8)

            if n_rows >= 2:
                for p in ht.cell(1, 0).paragraphs:
                    p.clear()
                for p in ht.cell(1, 1).paragraphs:
                    p.clear()
                run = ht.cell(1, 1).paragraphs[0].add_run("Progetto {{PROGETTO}}")
                run.font.size = Pt(8)
                for p in ht.cell(1, 2).paragraphs:
                    p.clear()
                run = ht.cell(1, 2).paragraphs[0].add_run("{{TIPO_DOCUMENTO}}")
                run.font.size = Pt(8)
                for p in ht.cell(1, 3).paragraphs:
                    p.clear()
                run = ht.cell(1, 3).paragraphs[0].add_run("Versione {{VERSIONE}}")
                run.font.size = Pt(8)

            if n_rows >= 3:
                for p in ht.cell(2, 0).paragraphs:
                    p.clear()
                # Cell 2,1 and 2,2 are merged (gridSpan=2) — write only to 2,1
                for p in ht.cell(2, 1).paragraphs:
                    p.clear()
                run = ht.cell(2, 1).paragraphs[0].add_run("{{NOME_DOCUMENTO}}")
                run.font.size = Pt(8)
                # Skip cell 2,2 (same merged cell as 2,1)
                for p in ht.cell(2, 3).paragraphs:
                    p.clear()
                run = ht.cell(2, 3).paragraphs[0].add_run("{{DATA}}")
                run.font.size = Pt(8)

# ── 5. Pulisci footer ───────────────────────────────────────────────────
for section in doc.sections:
    footer = section.footer
    if not footer:
        continue
    # Il footer contiene una tabella con:
    #   Col 0: nome file originale (da sostituire con placeholder)
    #   Col 1: Pag. X / Y (campo PAGE/NUMPAGES — preservare)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for ft in footer.tables:
        # Prima cella: contiene il nome del file — rimpiazza con placeholder
        cell0 = ft.cell(0, 0)
        for p in cell0.paragraphs:
            p.clear()
        run = cell0.paragraphs[0].add_run("{{NOME_DOCUMENTO}}")
        run.font.size = Pt(9)

doc.save(str(DST))
print(f"Template salvato: {DST}")
