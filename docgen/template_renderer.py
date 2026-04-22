"""Template Renderer — converte Markdown in DOCX usando il template aziendale.

Apre il template DOCX (con copertina, header, stili aziendali), sostituisce
i placeholder, poi parsa il Markdown e lo inserisce come contenuto del documento.
"""

from __future__ import annotations

import base64
import os
import re
import struct
import subprocess
import tempfile
import urllib.request
from datetime import datetime
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from rich.console import Console

console = Console()

# ── Costanti stile (allineate al template aziendale) ─────────────────────

COLOR_HEADING = RGBColor(0x2F, 0x54, 0x96)      # Blu #2F5496 (dal template)
COLOR_TABLE_HEADER = RGBColor(0x2F, 0x54, 0x96)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_CODE_INLINE = RGBColor(0xC7, 0x25, 0x4E)
COLOR_CODE_BG = "F5F5F5"

FONT_CODE = "Consolas"
SIZE_CODE = Pt(9)

# Dimensioni heading dal template: H1=16pt, H2=13pt
SIZE_H1 = Pt(16)
SIZE_H2 = Pt(13)
SIZE_H3 = Pt(11)
SIZE_H4 = Pt(10)

# Default template path
DEFAULT_TEMPLATE = Path(__file__).parent.parent / "templates" / "template_aziendale.docx"


# ── Helper functions ─────────────────────────────────────────────────────

def _replace_placeholder_in_runs(element, placeholder: str, value: str) -> None:
    """Sostituisce un placeholder nei run di tutti i paragrafi di un elemento."""
    for p in element.paragraphs:
        full_text = p.text
        if placeholder in full_text:
            # Ricostruisci il testo mantenendo la formattazione del primo run
            for run in p.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
                    return
            # Se il placeholder è spezzato tra run, ricostruisci
            p.clear()
            new_run = p.add_run(full_text.replace(placeholder, value))


def _replace_placeholder_in_table(table, placeholder: str, value: str) -> None:
    """Sostituisce un placeholder in tutte le celle di una tabella."""
    for row in table.rows:
        for cell in row.cells:
            _replace_placeholder_in_runs(cell, placeholder, value)


def _add_inline_formatting(paragraph, text: str) -> None:
    """Aggiunge testo con formattazione inline (bold, italic, code).

    Ordine di parsing: inline code (`...`), bold (**...**), italic (*...*).
    Il pattern underscore _..._ è ignorato per evitare falsi positivi con
    nomi variabili (es. codice_fiscale).
    """
    # Pattern robusto: inline code ha priorità massima
    pattern = re.compile(
        r'`([^`]+)`'            # inline code
        r'|\*\*(.+?)\*\*'      # bold (non-greedy)
        r'|\*(.+?)\*'          # italic (non-greedy)
        r'|([^`*]+)'           # testo normale (tutto tranne backtick e asterischi)
        r'|([`*])'             # caratteri residui singoli (backtick/asterisco solitari)
    )
    for match in pattern.finditer(text):
        if match.group(1):  # inline code
            run = paragraph.add_run(match.group(1))
            run.font.name = FONT_CODE
            run.font.size = SIZE_CODE
            # Colore code più discreto: grigio scuro anziché rosa
            run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
        elif match.group(2):  # bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # italic
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # testo normale
            run = paragraph.add_run(match.group(4))
        elif match.group(5):  # carattere residuo
            run = paragraph.add_run(match.group(5))


def _set_cell_shading(cell, color_hex: str) -> None:
    """Imposta sfondo a una cella."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_table_borders(table) -> None:
    """Imposta bordi 'Tutti i bordi' (single) sulla tabella."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _add_code_block(doc, code_lines: list[str]) -> None:
    """Inserisce un blocco di codice in una tabella 1x1 con sfondo grigio.

    Usa una tabella a cella singola per isolare il blocco codice e garantire
    allineamento a sinistra con spazi preservati (no giustificazione Word).
    """
    if not any(line.strip() for line in code_lines):
        return
    # Tabella 1x1 per contenere il codice — evita giustificazione Word
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    # Bordo sottile grigio
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
        table._tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="D9D9D9"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="D9D9D9"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="D9D9D9"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="D9D9D9"/>'
        '</w:tblBorders>'
    )
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(borders)
    # Sfondo grigio chiaro sulla cella
    _set_cell_shading(cell, COLOR_CODE_BG)
    # Usa il paragrafo esistente nella cella
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for idx, code_line in enumerate(code_lines):
        if idx > 0:
            br_run = p.add_run()
            br_run.font.name = FONT_CODE
            br_run.font.size = SIZE_CODE
            br_run.add_break()
        # Preserva spazi iniziali: sostituisci spazi con non-breaking spaces
        if code_line:
            leading = len(code_line) - len(code_line.lstrip(' '))
            display = '\u00A0' * leading + code_line[leading:]
        else:
            display = " "
        run = p.add_run(display)
        run.font.name = FONT_CODE
        run.font.size = SIZE_CODE


def _render_mermaid(doc, mermaid_code: str) -> bool:
    """Tenta di renderizzare un diagramma Mermaid come immagine nel documento.

    Strategia: mmdc locale → mermaid.ink API → fallback False.
    Scala automaticamente in base all'aspect ratio (max 15cm largo, 20cm alto).
    """
    MAX_WIDTH = Cm(15)
    MAX_HEIGHT = Cm(20)

    def _add_scaled_picture(doc, png_path):
        """Aggiunge immagine scalata al vincolo più restrittivo (larghezza o altezza)."""
        w_px, h_px = _get_image_dimensions(png_path)
        if w_px and h_px:
            aspect = w_px / h_px
            # Calcola dimensioni con entrambi i vincoli
            w_by_width = MAX_WIDTH
            h_by_width = int(MAX_WIDTH / aspect)
            w_by_height = int(MAX_HEIGHT * aspect)
            h_by_height = MAX_HEIGHT
            # Scegli il vincolo più restrittivo
            if h_by_width <= MAX_HEIGHT:
                doc.add_picture(png_path, width=w_by_width)
            else:
                doc.add_picture(png_path, height=h_by_height)
        else:
            doc.add_picture(png_path, width=MAX_WIDTH)

    # 1. Prova mmdc (mermaid CLI) se installato
    try:
        with tempfile.NamedTemporaryFile(
            suffix='.mmd', delete=False, mode='w', encoding='utf-8'
        ) as f:
            f.write(mermaid_code)
            mmd_path = f.name
        png_path = mmd_path.replace('.mmd', '.png')
        result = subprocess.run(
            ['mmdc', '-i', mmd_path, '-o', png_path, '-b', 'white', '-s', '2'],
            capture_output=True, timeout=30
        )
        if result.returncode == 0 and os.path.exists(png_path):
            _add_scaled_picture(doc, png_path)
            os.unlink(mmd_path)
            os.unlink(png_path)
            return True
        os.unlink(mmd_path)
        if os.path.exists(png_path):
            os.unlink(png_path)
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        pass

    # 2. Prova mermaid.ink API (servizio pubblico)
    try:
        encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('ascii')
        url = f"https://mermaid.ink/img/{encoded}"
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=20) as resp:
            img_data = resp.read()
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            f.write(img_data)
            png_path = f.name
        _add_scaled_picture(doc, png_path)
        os.unlink(png_path)
        return True
    except Exception:
        pass

    return False


def _get_image_dimensions(img_path: str) -> tuple[int | None, int | None]:
    """Legge larghezza e altezza da un file PNG o JPEG."""
    try:
        with open(img_path, 'rb') as f:
            header = f.read(32)

            # PNG: IHDR chunk a offset 16
            if header[:8] == b'\x89PNG\r\n\x1a\n':
                w, h = struct.unpack('>II', header[16:24])
                return w, h

            # JPEG: cerca il marker SOF0/SOF2 per le dimensioni
            if header[:2] == b'\xff\xd8':
                f.seek(0)
                data = f.read()
                i = 2
                while i < len(data) - 9:
                    if data[i] != 0xFF:
                        i += 1
                        continue
                    marker = data[i + 1]
                    # SOF0 (0xC0) o SOF2 (0xC2) contengono le dimensioni
                    if marker in (0xC0, 0xC2):
                        h, w = struct.unpack('>HH', data[i + 5:i + 9])
                        return w, h
                    # Skip non-dimension markers
                    if marker == 0xD9 or marker == 0xDA:
                        break
                    seg_len = struct.unpack('>H', data[i + 2:i + 4])[0]
                    i += 2 + seg_len
    except (OSError, struct.error, IndexError):
        pass
    return None, None


def _parse_table(lines: list[str]) -> list[list[str]]:
    """Parsa una tabella Markdown e restituisce righe di celle."""
    rows: list[list[str]] = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells:
            rows.append(cells)
    return rows


# ── Renderer principale ─────────────────────────────────────────────────

def render_md_to_docx(
    markdown_path: str,
    output_path: str,
    template_path: str | None = None,
    metadata: dict[str, str] | None = None,
) -> str:
    """Converte un file Markdown in DOCX usando il template aziendale.

    Args:
        markdown_path: percorso al file .md da convertire
        output_path: percorso output .docx
        template_path: percorso al template .docx (default: templates/template_aziendale.docx)
        metadata: dizionario con valori per i placeholder del template
            Chiavi supportate:
            - INTESTAZIONE_ENTE
            - NOME_DOCUMENTO
            - VERSIONE
            - STATO
            - DATA
            - CLIENTE
            - PROGETTO
            - REDATTO_DA
            - APPROVATO_DA
            - VERIFICATO_DA
            - TIPO_DOCUMENTO

    Returns:
        percorso del file .docx generato
    """
    template = Path(template_path) if template_path else DEFAULT_TEMPLATE
    if not template.exists():
        raise FileNotFoundError(f"Template non trovato: {template}")

    md_path = Path(markdown_path)
    if not md_path.exists():
        raise FileNotFoundError(f"File Markdown non trovato: {md_path}")

    markdown_text = md_path.read_text(encoding="utf-8")

    # Default metadata
    meta = {
        "INTESTAZIONE_ENTE": "XXXX",
        "NOME_DOCUMENTO": md_path.stem,
        "VERSIONE": "1.0",
        "STATO": "Bozza",
        "DATA": datetime.now().strftime("%d/%m/%Y"),
        "CLIENTE": "XXXX",
        "PROGETTO": "XXXX",
        "REDATTO_DA": "DocGen (generazione automatica)",
        "APPROVATO_DA": "XXXX",
        "VERIFICATO_DA": "XXXX",
        "TIPO_DOCUMENTO": _detect_doc_type(md_path.stem),
    }
    if metadata:
        meta.update(metadata)

    # Apri template
    doc = Document(str(template))

    # ── 1. Sostituisci placeholder nella copertina e metadati ────────
    for table in doc.tables:
        for placeholder_key, value in meta.items():
            _replace_placeholder_in_table(table, "{{" + placeholder_key + "}}", value)

    # ── 2. Sostituisci placeholder nell'header ───────────────────────
    for section in doc.sections:
        if section.header and section.header.tables:
            for ht in section.header.tables:
                for placeholder_key, value in meta.items():
                    _replace_placeholder_in_table(ht, "{{" + placeholder_key + "}}", value)

    # ── 2b. Sostituisci placeholder nel footer ──────────────────────
    for section in doc.sections:
        if section.footer:
            for ft in section.footer.tables:
                for placeholder_key, value in meta.items():
                    _replace_placeholder_in_table(ft, "{{" + placeholder_key + "}}", value)

    # ── 3. Inserisci il contenuto Markdown ───────────────────────────
    _insert_markdown_content(doc, markdown_text)

    # ── 4. Salva ─────────────────────────────────────────────────────
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    return str(out)


def _detect_doc_type(filename: str) -> str:
    """Rileva il tipo di documento dal nome file."""
    lower = filename.lower()
    if "funzionale" in lower:
        return "Specifica Funzionale"
    elif "tecnic" in lower:
        return "Specifica Tecnica"
    elif "architettura" in lower:
        return "Architettura di Sistema"
    return "Documentazione"


def _insert_markdown_content(doc: Document, markdown_text: str) -> None:
    """Parsa il Markdown e lo inserisce nel documento dopo le tabelle di copertina."""
    body = doc.element.body

    # Trova il punto di inserimento: dopo l'ultimo paragrafo vuoto (che è il segnaposto)
    # Aggiungiamo i paragrafi al documento normalmente — python-docx li appende alla fine
    # Il paragrafo vuoto del template fa da separatore dopo le tabelle di copertina

    # Page break dopo la copertina
    pb = doc.add_paragraph()
    run = pb.add_run()
    run.add_break(WD_BREAK.PAGE)

    # Parse Markdown line by line
    lines = markdown_text.split("\n")
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    in_mermaid = False

    while i < len(lines):
        line = lines[i]

        # ── Code block / Mermaid ─────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                in_mermaid = "mermaid" in line.strip().lower()
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                if not in_mermaid:
                    _add_code_block(doc, code_lines)
                else:
                    # Mermaid: prova rendering immagine, fallback a codice sorgente
                    mermaid_src = "\n".join(code_lines)
                    if not _render_mermaid(doc, mermaid_src):
                        # Fallback: mostra sorgente Mermaid come blocco codice
                        p = doc.add_paragraph()
                        run = p.add_run("[Diagramma Mermaid]")
                        run.italic = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                        _add_code_block(doc, code_lines)
                in_mermaid = False
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # ── Riga vuota ───────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Heading ──────────────────────────────────────────────────
        heading_match = re.match(r'^(#{1,4})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)

            # Usa gli stili Word standard che ereditano dal template
            style_name = f"Heading {level}" if level <= 3 else "Heading 3"
            try:
                p = doc.add_paragraph(text, style=style_name)
            except KeyError:
                # Fallback se lo stile non esiste
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                sizes = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3, 4: SIZE_H4}
                run.font.size = sizes.get(level, SIZE_H4)
                run.font.color.rgb = COLOR_HEADING
            i += 1
            continue

        # ── Tabella ──────────────────────────────────────────────────
        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            rows = _parse_table(table_lines)
            if not rows:
                continue

            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_borders(table)

            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < num_cols:
                        cell = table.cell(row_idx, col_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _add_inline_formatting(p, cell_text)
                        p.paragraph_format.space_after = Pt(2)

                        if row_idx == 0:
                            _set_cell_shading(cell, "2F5496")
                            for run in p.runs:
                                run.bold = True
                                run.font.color.rgb = COLOR_WHITE
                                run.font.size = Pt(10)
            continue

        # ── Lista puntata ────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*+]\s+(.+)', stripped)
        if bullet_match:
            text = bullet_match.group(2)
            p = doc.add_paragraph(style="List Paragraph")
            _add_inline_formatting(p, text)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # ── Lista numerata ───────────────────────────────────────────
        numbered_match = re.match(r'^(\s*)\d+\.\s+(.+)', stripped)
        if numbered_match:
            text = numbered_match.group(2)
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph(style="List Paragraph")
            _add_inline_formatting(p, text)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        # ── Separatore (hr) ──────────────────────────────────────────
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            # Linea sottile tramite bordo inferiore — aspetto professionale
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            border_xml = (
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
                '</w:pBdr>'
            )
            pPr.append(parse_xml(border_xml))
            i += 1
            continue

        # ── Paragrafo normale ────────────────────────────────────────
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_inline_formatting(p, stripped)
        i += 1


# ── Funzione batch per convertire più file ───────────────────────────────

def render_batch(
    md_files: list[str],
    output_dir: str,
    template_path: str | None = None,
    metadata: dict[str, str] | None = None,
) -> list[str]:
    """Converte una lista di file .md in .docx.

    Returns:
        lista dei percorsi .docx generati
    """
    generated: list[str] = []
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    for md_file in md_files:
        md_path = Path(md_file)
        if not md_path.exists() or md_path.suffix != ".md":
            console.print(f"  [yellow]⚠[/yellow] Saltato: {md_path}")
            continue

        docx_name = md_path.stem + ".docx"
        docx_path = out_dir / docx_name

        try:
            result = render_md_to_docx(
                str(md_path),
                str(docx_path),
                template_path=template_path,
                metadata=metadata,
            )
            generated.append(result)
            console.print(f"  [green]✓[/green] {docx_name}")
        except Exception as e:
            console.print(f"  [red]✗[/red] {docx_name}: {e}")

    return generated
