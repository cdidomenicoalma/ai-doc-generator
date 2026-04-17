"""Renderer — converte Markdown in DOCX con formattazione professionale."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from rich.console import Console

console = Console()

# ── Costanti stile ───────────────────────────────────────────────────────────

COLOR_HEADING = RGBColor(0x1A, 0x3C, 0x6E)     # Blu scuro
COLOR_CODE_INLINE = RGBColor(0xC7, 0x25, 0x4E)  # Rosa per inline code
COLOR_CODE_BG = RGBColor(0xF5, 0xF5, 0xF5)      # Sfondo grigio chiaro
COLOR_TABLE_HEADER = RGBColor(0x1A, 0x3C, 0x6E)  # Blu scuro header tabella
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT_BODY = "Calibri"
FONT_CODE = "Consolas"
SIZE_BODY = Pt(11)
SIZE_CODE = Pt(9)
SIZE_H1 = Pt(20)
SIZE_H2 = Pt(16)
SIZE_H3 = Pt(13)
SIZE_H4 = Pt(11)


# ── Funzioni helper ──────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str) -> None:
    """Imposta sfondo a una cella."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_inline_formatting(paragraph, text: str) -> None:
    """Aggiunge testo con formattazione inline (bold, italic, code)."""
    # Pattern per bold, italic, inline code
    # Ordine: code (`...`), bold (**...**), italic (*...* o _..._)
    pattern = re.compile(
        r'`([^`]+)`'           # inline code
        r'|\*\*([^*]+)\*\*'    # bold
        r'|\*([^*]+)\*'        # italic
        r'|_([^_]+)_'          # italic (underscore)
        r'|([^`*_]+)'          # testo normale
    )

    for match in pattern.finditer(text):
        if match.group(1):  # inline code
            run = paragraph.add_run(match.group(1))
            run.font.name = FONT_CODE
            run.font.size = SIZE_CODE
            run.font.color.rgb = COLOR_CODE_INLINE
        elif match.group(2):  # bold
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3):  # italic *
            run = paragraph.add_run(match.group(3))
            run.italic = True
        elif match.group(4):  # italic _
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # testo normale
            run = paragraph.add_run(match.group(5))


def _add_page_number(section) -> None:
    """Aggiunge numerazione pagine nel footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instr_text)

    run3 = paragraph.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fld_char_end)


# ── Renderer principale ─────────────────────────────────────────────────────

def _create_cover_page(doc: Document, title: str, project_name: str) -> None:
    """Crea la pagina di copertina."""
    # Spazio vuoto sopra
    for _ in range(6):
        doc.add_paragraph("")

    # Titolo documento
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_HEADING
    run.bold = True
    run.font.name = FONT_BODY

    # Linea separatrice
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    run.font.color.rgb = COLOR_HEADING

    # Nome progetto
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project_name)
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_HEADING
    run.font.name = FONT_BODY

    # Data
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%d/%m/%Y"))
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = FONT_BODY

    # Generato da
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento generato automaticamente da DocGen")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    run.font.name = FONT_BODY

    # Page break dopo copertina
    doc.add_page_break()


def _parse_table(lines: list[str]) -> list[list[str]]:
    """Parsa una tabella Markdown e restituisce righe di celle."""
    rows: list[list[str]] = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        # Salta riga separatore (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells:
            rows.append(cells)
    return rows


def markdown_to_docx(
    markdown_text: str,
    output_path: str,
    title: str = "Documento",
    project_name: str = "",
) -> None:
    """Converte testo Markdown in un documento DOCX formattato."""
    doc = Document()

    # Stile default del documento
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = SIZE_BODY
    style.paragraph_format.space_after = Pt(6)

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Copertina
    _create_cover_page(doc, title, project_name)

    # Numerazione pagine
    for section in doc.sections:
        _add_page_number(section)

    # Parse Markdown linea per linea
    lines = markdown_text.split("\n")
    i = 0
    in_code_block = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        # ── Code block ───────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                # Fine code block — renderizza
                in_code_block = False
                code_text = "\n".join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run(code_text)
                    run.font.name = FONT_CODE
                    run.font.size = SIZE_CODE
                    # Sfondo grigio tramite shading
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'
                    )
                    p._p.get_or_add_pPr().append(shading)
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # ── Riga vuota ───────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Heading ──────────────────────────────────────────────────
        heading_match = re.match(r'^(#{1,4})\s+(.+)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            sizes = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3, 4: SIZE_H4}

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12) if level <= 2 else Pt(8)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.bold = True
            run.font.size = sizes.get(level, SIZE_H4)
            run.font.color.rgb = COLOR_HEADING
            run.font.name = FONT_BODY
            i += 1
            continue

        # ── Tabella ──────────────────────────────────────────────────
        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1

            rows = _parse_table(table_lines)
            if not rows:
                continue

            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < num_cols:
                        cell = table.cell(row_idx, col_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _add_inline_formatting(p, cell_text)
                        p.paragraph_format.space_after = Pt(2)

                        # Header row: sfondo blu, testo bianco bold
                        if row_idx == 0:
                            _set_cell_shading(cell, "1A3C6E")
                            for run in p.runs:
                                run.bold = True
                                run.font.color.rgb = COLOR_WHITE
                                run.font.size = Pt(10)

            continue

        # ── Lista puntata ────────────────────────────────────────────
        bullet_match = re.match(r'^(\s*)[-*+]\s+(.+)', stripped)
        if bullet_match:
            text = bullet_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, text)
            i += 1
            continue

        # ── Lista numerata ───────────────────────────────────────────
        numbered_match = re.match(r'^(\s*)\d+\.\s+(.+)', stripped)
        if numbered_match:
            text = numbered_match.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, text)
            i += 1
            continue

        # ── Separatore ───────────────────────────────────────────────
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            # Aggiungi una linea sottile
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # ── Paragrafo normale ────────────────────────────────────────
        p = doc.add_paragraph()
        _add_inline_formatting(p, stripped)
        i += 1

    # Salva
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def render_documents(
    functional_md: str,
    technical_md: str,
    output_dir: str,
    project_name: str,
    output_format: str = "all",
) -> list[str]:
    """Renderizza i documenti nei formati richiesti.
    
    Ritorna la lista dei file generati.
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    generated: list[str] = []

    func_base = f"DOC_FUNZIONALE_{timestamp}"
    tech_base = f"DOC_TECNICA_{timestamp}"

    # Markdown
    if output_format in ("all", "md"):
        func_md_path = output_path / f"{func_base}.md"
        tech_md_path = output_path / f"{tech_base}.md"

        func_md_path.write_text(functional_md, encoding="utf-8")
        tech_md_path.write_text(technical_md, encoding="utf-8")

        generated.extend([str(func_md_path), str(tech_md_path)])
        console.print(f"  [green]✓[/green] {func_md_path.name}")
        console.print(f"  [green]✓[/green] {tech_md_path.name}")

    # DOCX
    if output_format in ("all", "docx"):
        func_docx_path = output_path / f"{func_base}.docx"
        tech_docx_path = output_path / f"{tech_base}.docx"

        console.print("  Rendering DOCX funzionale...")
        markdown_to_docx(
            functional_md,
            str(func_docx_path),
            title="Specifica Funzionale",
            project_name=project_name,
        )
        console.print(f"  [green]✓[/green] {func_docx_path.name}")

        console.print("  Rendering DOCX tecnica...")
        markdown_to_docx(
            technical_md,
            str(tech_docx_path),
            title="Specifica Tecnica",
            project_name=project_name,
        )
        console.print(f"  [green]✓[/green] {tech_docx_path.name}")

        generated.extend([str(func_docx_path), str(tech_docx_path)])

    return generated


def render_documents_hybrid(
    results: dict[str, tuple[str, str]],
    output_dir: str,
    project_name: str,
    output_format: str = "all",
) -> list[str]:
    """Renderizza i documenti della generazione ibrida.
    
    Ogni microservizio ha la propria sottodirectory.
    Il documento di architettura di sistema va nella root.
    
    Ritorna la lista dei file generati.
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    generated: list[str] = []

    for module_name, (func_md, tech_md) in sorted(results.items()):
        # Documento di architettura di sistema → root
        if module_name == "_architettura_sistema":
            arch_base = f"DOC_ARCHITETTURA_SISTEMA_{timestamp}"
            console.print(f"\n  [bold]Architettura di sistema[/bold]")

            if output_format in ("all", "md"):
                arch_md_path = output_path / f"{arch_base}.md"
                arch_md_path.write_text(func_md, encoding="utf-8")
                generated.append(str(arch_md_path))
                console.print(f"  [green]✓[/green] {arch_md_path.name}")

            if output_format in ("all", "docx"):
                arch_docx_path = output_path / f"{arch_base}.docx"
                console.print("  Rendering DOCX architettura...")
                markdown_to_docx(
                    func_md,
                    str(arch_docx_path),
                    title="Architettura di Sistema",
                    project_name=project_name,
                )
                console.print(f"  [green]✓[/green] {arch_docx_path.name}")
                generated.append(str(arch_docx_path))
            continue

        # Microservizio → sottodirectory
        module_path = output_path / module_name
        module_path.mkdir(parents=True, exist_ok=True)

        func_base = f"DOC_FUNZIONALE_{module_name}_{timestamp}"
        tech_base = f"DOC_TECNICA_{module_name}_{timestamp}"

        console.print(f"\n  [bold]{module_name}[/bold]")

        if output_format in ("all", "md"):
            func_md_path = module_path / f"{func_base}.md"
            tech_md_path = module_path / f"{tech_base}.md"

            func_md_path.write_text(func_md, encoding="utf-8")
            tech_md_path.write_text(tech_md, encoding="utf-8")

            generated.extend([str(func_md_path), str(tech_md_path)])
            console.print(f"  [green]✓[/green] {module_name}/{func_md_path.name}")
            console.print(f"  [green]✓[/green] {module_name}/{tech_md_path.name}")

        if output_format in ("all", "docx"):
            func_docx_path = module_path / f"{func_base}.docx"
            tech_docx_path = module_path / f"{tech_base}.docx"

            console.print(f"  Rendering DOCX funzionale {module_name}...")
            markdown_to_docx(
                func_md,
                str(func_docx_path),
                title=f"Specifica Funzionale — {module_name}",
                project_name=project_name,
            )
            console.print(f"  [green]✓[/green] {module_name}/{func_docx_path.name}")

            console.print(f"  Rendering DOCX tecnica {module_name}...")
            markdown_to_docx(
                tech_md,
                str(tech_docx_path),
                title=f"Specifica Tecnica — {module_name}",
                project_name=project_name,
            )
            console.print(f"  [green]✓[/green] {module_name}/{tech_docx_path.name}")

            generated.extend([str(func_docx_path), str(tech_docx_path)])

    return generated
